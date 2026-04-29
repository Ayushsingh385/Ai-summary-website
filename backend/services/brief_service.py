"""
Brief Generator Service
──────────────────────
Generates structured legal briefs from case documents.
Uses NLP to extract: Issues, Facts, Argument, Prayer.
Outputs DOCX with formal legal formatting.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

from services.legal_templates import (
    apply_template,
    _styled_paragraph,
    _add_section_heading,
    _add_horizontal_line,
    COLORS,
)
from services.nlp_service import summarize_text, extract_keywords


BRIEF_TYPES = {
    "memo": "Legal Memorandum",
    "brief": "Legal Brief",
    "opinion": "Legal Opinion",
    "summary": "Case Summary",
}


def generate_brief_docx(
    filename: str,
    original_text: str,
    summary: str,
    keywords: list,
    legal_analysis: dict = None,
    case_type: dict = None,
    brief_type: str = "memo",
    template: str = "general",
) -> bytes:
    """
    Generate a structured legal brief as DOCX.

    Sections:
        1. Header / Cover
        2. Issues Presented
        3. Facts of the Case
        4. Legal Analysis / Arguments
        5. Prayer / Relief Sought
        6. Supporting Authorities

    Returns raw DOCX bytes.
    """
    doc = Document()

    # Apply base template
    title = BRIEF_TYPES.get(brief_type, "Legal Brief")
    apply_template(doc, template, title=title)

    # ── Case Cover / Metadata ─────────────────────────────────
    meta_items = [
        ("Document", filename),
        ("Brief Type", BRIEF_TYPES.get(brief_type, brief_type.title())),
        ("Generated", datetime.now().strftime("%d %B %Y, %I:%M %p")),
    ]
    if case_type and case_type.get("primary_type"):
        meta_items.append(("Case Category", case_type["primary_type"]))

    for key, val in meta_items:
        p = doc.add_paragraph()
        r_k = p.add_run(f"{key}: ")
        r_k.font.name = "Arial"
        r_k.font.size = Pt(10)
        r_k.font.bold = True
        r_k.font.color.rgb = COLORS["muted"]
        r_v = p.add_run(val)
        r_v.font.name = "Arial"
        r_v.font.size = Pt(10)
        r_v.font.color.rgb = COLORS["body_text"]
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # ── Section 1: Issues Presented ────────────────────────────
    _add_section_heading(doc, "I. ISSUES PRESENTED", level=2,
                         color=COLORS.get({"zp_official": "zp_primary",
                                           "court_order": "court_primary",
                                           "general": "gen_primary"}.get(template, "gen_primary"),
                                           COLORS["gen_primary"]))

    issues = _extract_issues(original_text, legal_analysis)
    for i, issue in enumerate(issues, 1):
        _styled_paragraph(
            doc, f"{i}. {issue}",
            font_name="Times New Roman", font_size=11,
            color=COLORS["body_text"],
            space_after=Pt(6),
        )
    doc.add_paragraph()

    # ── Section 2: Facts of the Case ──────────────────────────
    _add_section_heading(doc, "II. FACTS OF THE CASE", level=2,
                         color=COLORS.get({"zp_official": "zp_primary",
                                           "court_order": "court_primary",
                                           "general": "gen_primary"}.get(template, "gen_primary"),
                                           COLORS["gen_primary"]))

    facts = _extract_facts(original_text, legal_analysis, summary)
    _styled_paragraph(
        doc, facts,
        font_name="Times New Roman", font_size=11,
        color=COLORS["body_text"],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=Pt(8),
    )
    doc.add_paragraph()

    # ── Section 3: Legal Analysis / Arguments ──────────────────
    _add_section_heading(doc, "III. LEGAL ANALYSIS", level=2,
                         color=COLORS.get({"zp_official": "zp_primary",
                                           "court_order": "court_primary",
                                           "general": "gen_primary"}.get(template, "gen_primary"),
                                           COLORS["gen_primary"]))

    analysis = _extract_analysis(original_text, legal_analysis, summary, keywords)
    _styled_paragraph(
        doc, analysis,
        font_name="Times New Roman", font_size=11,
        color=COLORS["body_text"],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=Pt(8),
    )
    doc.add_paragraph()

    # ── Section 4: Prayer / Relief Sought ─────────────────────
    _add_section_heading(doc, "IV. PRAYER / RELIEF SOUGHT", level=2,
                         color=COLORS.get({"zp_official": "zp_primary",
                                           "court_order": "court_primary",
                                           "general": "gen_primary"}.get(template, "gen_primary"),
                                           COLORS["gen_primary"]))

    prayers = _extract_prayer(original_text, legal_analysis)
    for prayer in prayers:
        _styled_paragraph(
            doc, f"• {prayer}",
            font_name="Times New Roman", font_size=11,
            color=COLORS["body_text"],
            space_after=Pt(4),
        )
    doc.add_paragraph()

    # ── Section 5: Supporting Authorities ──────────────────────
    _add_section_heading(doc, "V. SUPPORTING AUTHORITIES", level=2,
                         color=COLORS.get({"zp_official": "zp_primary",
                                           "court_order": "court_primary",
                                           "general": "gen_primary"}.get(template, "gen_primary"),
                                           COLORS["gen_primary"]))

    authorities = _extract_authorities(original_text, keywords)
    if authorities:
        for auth in authorities:
            _styled_paragraph(
                doc, f"• {auth}",
                font_name="Times New Roman", font_size=10, italic=True,
                color=COLORS["body_text"],
                space_after=Pt(3),
            )
    else:
        _styled_paragraph(
            doc, "No specific legal citations identified in the document.",
            font_name="Times New Roman", font_size=10, italic=True,
            color=COLORS["muted"],
            space_after=Pt(3),
        )

    doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────
    from services.legal_templates import add_footer_note
    add_footer_note(doc,
        note="Generated by Zilla Parishad AI Document Processor — Brief Generator",
        template_name=template)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────
# NLP helpers — extract each brief section
# ─────────────────────────────────────────────────────────────────

def _extract_issues(text: str, legal_analysis: dict = None) -> list:
    """Extract legal issues from document."""
    issues = []

    if legal_analysis and legal_analysis.get("legal_issues"):
        issues = legal_analysis["legal_issues"][:5]

    if not issues:
        # Fallback: use BART to generate issues
        prompt = (
            "From the following legal case text, identify the 3-5 main legal issues or questions "
            "presented. List each issue as a clear, specific question. Format: numbered list only.\n\n"
            f"Text:\n{text[:3000]}"
        )
        try:
            result = summarize_text(prompt, "short", "en")
            raw = result.get("summary", "")
            issues = [line.strip() for line in raw.split("\n") if line.strip() and (line[0].isdigit() or line.startswith("-"))][:5]
        except Exception:
            pass

    if not issues:
        issues = [
            "Whether the facts establish a prima facie case for the relief sought.",
            "What legal provisions are applicable to the matter.",
            "Whether procedural requirements have been duly complied with.",
        ]

    return issues


def _extract_facts(text: str, legal_analysis: dict = None, summary: str = "") -> str:
    """Extract a structured facts叙述."""
    if legal_analysis and legal_analysis.get("timeline"):
        timeline = legal_analysis["timeline"]
        if timeline:
            parts = []
            for event in timeline[:8]:
                date = event.get("date", "")
                desc = event.get("description", event.get("event", ""))
                if date:
                    parts.append(f"On {date}: {desc}")
                else:
                    parts.append(f"• {desc}")
            return "\n".join(parts)

    if summary:
        return summary

    return text[:1500] + ("..." if len(text) > 1500 else "")


def _extract_analysis(text: str, legal_analysis: dict = None,
                       summary: str = "", keywords: list = None) -> str:
    """Extract legal arguments and analysis."""
    if legal_analysis and legal_analysis.get("sections", {}).get("reasoning"):
        reasoning = legal_analysis["sections"]["reasoning"]
        if reasoning:
            return reasoning

    kw_list = [k.get("keyword", str(k)) if isinstance(k, dict) else str(k) for k in (keywords or [])]
    kw_str = ", ".join(kw_list[:10])

    prompt = (
        "Provide a concise legal analysis of the following case text. "
        "Identify the applicable law, the relevant facts, and how the law applies to the facts. "
        "Be analytical and precise.\n\n"
        f"Text:\n{text[:3000]}\n\nKeywords: {kw_str}"
    )
    try:
        result = summarize_text(prompt, "medium", "en")
        return result.get("summary", "")
    except Exception:
        return summary or "Analysis could not be generated from the provided text."


def _extract_prayer(text: str, legal_analysis: dict = None) -> list:
    """Extract the prayer/relief sought."""
    prayer_signals = [
        "prayer", "relief", "requested", "sought", "humbly", "implore",
        "granted", "order", "direction", "disposed of"
    ]
    text_lower = text.lower()

    sentences = []
    for i, sent in enumerate(text.split(".")):
        sent = sent.strip()
        if not sent:
            continue
        if any(signal in sent.lower() for signal in prayer_signals):
            sentences.append(sent)
        if len(sentences) >= 5:
            break

    if sentences:
        return sentences

    if legal_analysis and legal_analysis.get("sections", {}).get("order"):
        order_text = legal_analysis["sections"]["order"]
        return [order_text[:500]] if order_text else []

    return [
        "It is respectfully prayed that this Hon'ble Court may be pleased to pass appropriate orders.",
        "Costs of the proceedings may be awarded to the applicant.",
        "Such further and other relief as the Court deems fit may be granted.",
    ]


def _extract_authorities(text: str, keywords: list = None) -> list:
    """Extract referenced legal authorities."""
    authorities = []

    if keywords:
        law_keywords = [k.get("keyword", str(k)) for k in keywords
                        if isinstance(k, dict) and k.get("type") == "LAW"]
        authorities.extend(law_keywords[:5])

    # Regex find common statute references
    import re
    statutes = re.findall(
        r'(?:Section|Sec\.|Article|Art\.|Art\.)\s+\d+[\w\s,/-]*?(?:of|Under|under)\s+[A-Z][\w\s]+',
        text
    )
    for s in statutes[:5]:
        clean = s.strip()
        if clean not in authorities:
            authorities.append(clean)

    return authorities[:8]
