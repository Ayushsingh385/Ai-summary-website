"""
Legal Document Template Engine
─────────────────────────────
Provides reusable document templates for Zilla Parishad and legal exports.
Templates configure page layout, typography, header/footer, and structured
section builders that can be applied to any python-docx Document.

Available templates:
    • zp_official   — Zilla Parishad institutional format
    • court_order   — Formal court / tribunal order format
    • general       — Clean administrative/memo style
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re


# ────────────────────────────────────────────────────────────────
# Colour palette
# ────────────────────────────────────────────────────────────────
COLORS = {
    "zp_primary":    RGBColor(0x1A, 0x3C, 0x6E),   # Deep navy
    "zp_accent":     RGBColor(0x8B, 0x4C, 0x13),   # Official gold-brown
    "court_primary": RGBColor(0x2D, 0x2D, 0x2D),   # Near-black
    "court_accent":  RGBColor(0x6B, 0x21, 0x21),   # Dark maroon
    "gen_primary":   RGBColor(0x1F, 0x4E, 0x79),   # Professional blue
    "gen_accent":    RGBColor(0x44, 0x72, 0xC4),   # Accent blue
    "muted":         RGBColor(0x66, 0x66, 0x66),
    "light_muted":   RGBColor(0x99, 0x99, 0x99),
    "body_text":     RGBColor(0x1A, 0x1A, 0x1A),
}


# ────────────────────────────────────────────────────────────────
# Shared utilities
# ────────────────────────────────────────────────────────────────

def _set_page_layout(doc, top=2.0, bottom=2.0, left=2.54, right=2.54):
    """Set A4 page with given margins (cm)."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def _add_page_number(doc):
    """Add centred page number to the footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.style.font.size = Pt(8)
    para.style.font.color.rgb = COLORS["light_muted"]

    # Add "Page X" via field codes
    run = para.add_run("— Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["light_muted"]

    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = para.add_run()
    run1._r.append(fldChar1)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2 = para.add_run()
    run2._r.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3 = para.add_run()
    run3._r.append(fldChar2)

    run4 = para.add_run(" —")
    run4.font.size = Pt(8)
    run4.font.color.rgb = COLORS["light_muted"]


def _add_horizontal_line(doc, color=COLORS["muted"], thickness=0.75):
    """Insert a thin horizontal line via a paragraph bottom-border."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{int(thickness * 8)}" w:space="1" '
        f'    w:color="{color.hex_string if hasattr(color, "hex_string") else "666666"}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return para


def _styled_paragraph(doc, text, font_name="Times New Roman", font_size=11,
                       bold=False, italic=False, color=None, alignment=None,
                       space_after=Pt(6), space_before=Pt(0), line_spacing=1.15):
    """Add a paragraph with precise styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if alignment:
        para.alignment = alignment
    fmt = para.paragraph_format
    fmt.space_after = space_after
    fmt.space_before = space_before
    fmt.line_spacing = line_spacing
    return para


def _add_metadata_block(doc, items, color=COLORS["muted"]):
    """Add a compact key-value metadata block."""
    for key, value in items:
        para = doc.add_paragraph()
        run_key = para.add_run(f"{key}: ")
        run_key.font.name = "Arial"
        run_key.font.size = Pt(9)
        run_key.font.bold = True
        run_key.font.color.rgb = color
        run_val = para.add_run(str(value))
        run_val.font.name = "Arial"
        run_val.font.size = Pt(9)
        run_val.font.color.rgb = COLORS["body_text"]
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(0)


def _add_section_heading(doc, text, level=1, color=COLORS["gen_primary"],
                          font_name="Arial"):
    """Add a styled section heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        run.font.color.rgb = color
    return heading


def _shading_element(color_hex):
    """Create cell shading XML."""
    return parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )


def _style_table_cell(cell, text, bold=False, font_size=9, bg_color=None,
                       text_color=COLORS["body_text"], font_name="Arial"):
    """Style a table cell with text."""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = text_color
    if bg_color:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(_shading_element(bg_color))


# ────────────────────────────────────────────────────────────────
# Template: Zilla Parishad Official
# ────────────────────────────────────────────────────────────────

def apply_zp_official(doc, title="Case Summary Report", reference_no=None):
    """
    Apply Zilla Parishad official institutional formatting.
    Adds a government-style header band with institution name,
    reference number, and date.
    """
    _set_page_layout(doc, top=2.5, bottom=2.0, left=2.54, right=2.54)
    _add_page_number(doc)

    now = datetime.now()
    ref = reference_no or f"ZP/{now.strftime('%Y')}/{now.strftime('%m%d')}-{now.microsecond % 1000:03d}"

    # ── Institution header ──
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("जिल्हा परिषद  •  ZILLA PARISHAD")
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLORS["zp_primary"]
    header_para.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Administrative Document Processing System")
    run2.font.name = "Arial"
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = COLORS["zp_accent"]
    sub.paragraph_format.space_after = Pt(4)

    _add_horizontal_line(doc, COLORS["zp_primary"], 1.5)

    # ── Reference / Date bar ──
    ref_para = doc.add_paragraph()
    ref_run = ref_para.add_run(f"Ref. No: {ref}")
    ref_run.font.name = "Arial"
    ref_run.font.size = Pt(8)
    ref_run.font.color.rgb = COLORS["muted"]

    date_run = ref_para.add_run(f"      Date: {now.strftime('%d %B %Y')}")
    date_run.font.name = "Arial"
    date_run.font.size = Pt(8)
    date_run.font.color.rgb = COLORS["muted"]
    ref_para.paragraph_format.space_after = Pt(12)

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_para.add_run(title)
    t_run.font.name = "Arial"
    t_run.font.size = Pt(18)
    t_run.font.bold = True
    t_run.font.color.rgb = COLORS["zp_primary"]
    title_para.paragraph_format.space_after = Pt(6)

    _add_horizontal_line(doc, COLORS["zp_accent"], 0.5)
    doc.add_paragraph()  # spacer

    return doc


# ────────────────────────────────────────────────────────────────
# Template: Court Order
# ────────────────────────────────────────────────────────────────

def apply_court_order(doc, title="Case Summary Report",
                       case_number=None, court_name=None):
    """
    Apply formal court/tribunal order formatting.
    Includes court name, case number, date, and formal structure.
    """
    _set_page_layout(doc, top=3.0, bottom=2.5, left=3.0, right=2.54)
    _add_page_number(doc)

    now = datetime.now()
    c_num = case_number or f"WP/{now.strftime('%Y')}/{now.microsecond % 10000:04d}"
    c_name = court_name or "HIGH COURT OF JUDICATURE"

    # ── Court header ──
    court_para = doc.add_paragraph()
    court_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = court_para.add_run(c_name)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLORS["court_primary"]
    court_para.paragraph_format.space_after = Pt(4)

    # ── Case number ──
    case_para = doc.add_paragraph()
    case_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = case_para.add_run(f"Case No: {c_num}")
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(11)
    cr.font.color.rgb = COLORS["court_accent"]
    case_para.paragraph_format.space_after = Pt(2)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run(f"Date of Order: {now.strftime('%d/%m/%Y')}")
    dr.font.name = "Times New Roman"
    dr.font.size = Pt(10)
    dr.font.color.rgb = COLORS["muted"]
    date_para.paragraph_format.space_after = Pt(8)

    _add_horizontal_line(doc, COLORS["court_primary"], 2.0)
    doc.add_paragraph()

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_para.add_run(title)
    t_run.font.name = "Times New Roman"
    t_run.font.size = Pt(16)
    t_run.font.bold = True
    t_run.font.color.rgb = COLORS["court_primary"]
    t_run.font.underline = True
    title_para.paragraph_format.space_after = Pt(12)

    return doc


# ────────────────────────────────────────────────────────────────
# Template: General Administrative
# ────────────────────────────────────────────────────────────────

def apply_general(doc, title="Case Summary Report"):
    """
    Apply a clean, professional administrative/memo style.
    Minimalist design with blue accent tones.
    """
    _set_page_layout(doc, top=2.0, bottom=2.0, left=2.54, right=2.54)
    _add_page_number(doc)

    now = datetime.now()

    # ── Accent bar (thin top line) ──
    bar = doc.add_paragraph()
    bar.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bar_pPr = bar._p.get_or_add_pPr()
    bar_bdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="24" w:space="1" w:color="1F4E79"/>'
        f'</w:pBdr>'
    )
    bar_pPr.append(bar_bdr)
    bar.paragraph_format.space_after = Pt(8)

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_run = title_para.add_run(title)
    t_run.font.name = "Arial"
    t_run.font.size = Pt(20)
    t_run.font.bold = True
    t_run.font.color.rgb = COLORS["gen_primary"]
    title_para.paragraph_format.space_after = Pt(4)

    # ── Date ──
    d_para = doc.add_paragraph()
    d_run = d_para.add_run(f"Generated: {now.strftime('%d %B %Y, %I:%M %p')}")
    d_run.font.name = "Arial"
    d_run.font.size = Pt(9)
    d_run.font.italic = True
    d_run.font.color.rgb = COLORS["light_muted"]
    d_para.paragraph_format.space_after = Pt(12)

    _add_horizontal_line(doc, COLORS["gen_accent"], 0.5)
    doc.add_paragraph()

    return doc


# ────────────────────────────────────────────────────────────────
# Template dispatcher
# ────────────────────────────────────────────────────────────────

TEMPLATES = {
    "zp_official": apply_zp_official,
    "court_order": apply_court_order,
    "general":     apply_general,
}


def apply_template(doc, template_name, title="Case Summary Report", **kwargs):
    """
    Apply a named template to a Document.
    Returns the (modified) Document.

    If template_name is None or unrecognised, returns doc unchanged.
    """
    fn = TEMPLATES.get(template_name)
    if fn:
        return fn(doc, title=title, **kwargs)
    return doc


# ────────────────────────────────────────────────────────────────
# Shared content builders (used by download_service)
# ────────────────────────────────────────────────────────────────

def add_statistics_section(doc, items, template_name=None):
    """
    Add a compact statistics box.
    items: list of (label, value) tuples.
    """
    primary = COLORS.get(
        {"zp_official": "zp_primary", "court_order": "court_primary",
         "general": "gen_primary"}.get(template_name, "gen_primary"),
        COLORS["gen_primary"]
    )

    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, (label, value) in enumerate(items):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r_val = p.add_run(f"{value}\n")
        r_val.font.name = "Arial"
        r_val.font.size = Pt(14)
        r_val.font.bold = True
        r_val.font.color.rgb = primary

        r_lbl = p.add_run(label)
        r_lbl.font.name = "Arial"
        r_lbl.font.size = Pt(8)
        r_lbl.font.color.rgb = COLORS["muted"]

        # Light background
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(_shading_element("F5F5F5"))

    doc.add_paragraph()  # spacer


def add_body_text(doc, text, template_name=None):
    """Add the main body text, split into paragraphs, with proper styling."""
    font = "Times New Roman" if template_name == "court_order" else "Times New Roman"
    paragraphs = re.split(r'\n{2,}', text.strip())

    for chunk in paragraphs:
        chunk = chunk.strip()
        if not chunk:
            continue
        para = _styled_paragraph(
            doc, chunk,
            font_name=font,
            font_size=11,
            color=COLORS["body_text"],
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=Pt(8),
            line_spacing=1.15,
        )


def add_keywords_section(doc, keywords, template_name=None):
    """Add a styled keywords section using a table layout."""
    if not keywords:
        return

    primary = COLORS.get(
        {"zp_official": "zp_primary", "court_order": "court_primary",
         "general": "gen_primary"}.get(template_name, "gen_primary"),
        COLORS["gen_primary"]
    )

    _add_section_heading(doc, "Keywords & Legal Entities", level=2, color=primary)

    # Build a 3-column table for keywords
    cols = 3
    rows_needed = (len(keywords[:15]) + cols - 1) // cols
    table = doc.add_table(rows=rows_needed, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, kw in enumerate(keywords[:15]):
        row = i // cols
        col = i % cols
        keyword_text = kw.get("keyword", kw) if isinstance(kw, dict) else str(kw)
        keyword_type = kw.get("type", "") if isinstance(kw, dict) else ""

        cell = table.cell(row, col)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"• {keyword_text}")
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = COLORS["body_text"]
        if keyword_type:
            r2 = p.add_run(f"  ({keyword_type})")
            r2.font.name = "Arial"
            r2.font.size = Pt(8)
            r2.font.italic = True
            r2.font.color.rgb = COLORS["muted"]

    doc.add_paragraph()


def add_comparison_tables(doc, similarities, differences, shared_blocks,
                           shared_topics, unique_doc1, unique_doc2,
                           template_name=None):
    """Add comparison result tables for the comparison DOCX export."""
    primary = COLORS.get(
        {"zp_official": "zp_primary", "court_order": "court_primary",
         "general": "gen_primary"}.get(template_name, "gen_primary"),
        COLORS["gen_primary"]
    )

    # ── Identical Content Blocks ──
    if shared_blocks:
        _add_section_heading(doc, "Identical Content Blocks", level=2, color=primary)
        for i, block in enumerate(shared_blocks, 1):
            para = _styled_paragraph(
                doc, f"{i}. \"{block}\"",
                font_name="Times New Roman", font_size=10, italic=True,
                color=COLORS["body_text"],
                space_after=Pt(4),
            )

        doc.add_paragraph()

    # ── Similarities ──
    if similarities:
        _add_section_heading(doc, "Similarities Found", level=2, color=RGBColor(0x22, 0x8B, 0x22))

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Header row
        _style_table_cell(table.cell(0, 0), "Category", bold=True, bg_color="E8F5E9",
                           text_color=RGBColor(0x22, 0x8B, 0x22))
        _style_table_cell(table.cell(0, 1), "Shared Items", bold=True, bg_color="E8F5E9",
                           text_color=RGBColor(0x22, 0x8B, 0x22))

        for sim in similarities:
            cat = sim.get("category", "General") if isinstance(sim, dict) else str(sim)
            items = sim.get("items", []) if isinstance(sim, dict) else []
            row = table.add_row()
            _style_table_cell(row.cells[0], cat, bold=True)
            _style_table_cell(row.cells[1], ", ".join(str(x) for x in items))

        doc.add_paragraph()

    # ── Differences ──
    if differences:
        _add_section_heading(doc, "Differences Found", level=2, color=RGBColor(0xE6, 0x6A, 0x17))

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _style_table_cell(table.cell(0, 0), "Category", bold=True, bg_color="FFF3E0",
                           text_color=RGBColor(0xE6, 0x6A, 0x17))
        _style_table_cell(table.cell(0, 1), "Only in Document 1", bold=True, bg_color="FFF3E0",
                           text_color=RGBColor(0xE6, 0x6A, 0x17))
        _style_table_cell(table.cell(0, 2), "Only in Document 2", bold=True, bg_color="FFF3E0",
                           text_color=RGBColor(0xE6, 0x6A, 0x17))

        for diff in differences:
            cat = diff.get("category", "General") if isinstance(diff, dict) else str(diff)
            doc1 = diff.get("only_in_doc1", []) if isinstance(diff, dict) else []
            doc2 = diff.get("only_in_doc2", []) if isinstance(diff, dict) else []
            row = table.add_row()
            _style_table_cell(row.cells[0], cat, bold=True)
            _style_table_cell(row.cells[1], ", ".join(str(x) for x in doc1) or "—")
            _style_table_cell(row.cells[2], ", ".join(str(x) for x in doc2) or "—")

        doc.add_paragraph()

    # ── Shared Topics ──
    if shared_topics:
        _add_section_heading(doc, "Shared Topics", level=2, color=primary)
        topics_text = "  •  ".join(str(t) for t in shared_topics)
        _styled_paragraph(doc, topics_text, font_name="Arial", font_size=10,
                          color=COLORS["body_text"])
        doc.add_paragraph()

    # ── Unique Topics ──
    if unique_doc1 or unique_doc2:
        _add_section_heading(doc, "Unique Topics", level=2, color=RGBColor(0xE6, 0x6A, 0x17))

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _style_table_cell(table.cell(0, 0), "Only in Document 1", bold=True, bg_color="F3E5F5",
                           text_color=RGBColor(0x8B, 0x5C, 0xF6))
        _style_table_cell(table.cell(0, 1), "Only in Document 2", bold=True, bg_color="E3F2FD",
                           text_color=RGBColor(0x38, 0xBD, 0xF8))

        max_rows = max(len(unique_doc1 or []), len(unique_doc2 or []))
        for i in range(max_rows):
            row = table.add_row()
            _style_table_cell(row.cells[0], str(unique_doc1[i]) if i < len(unique_doc1 or []) else "")
            _style_table_cell(row.cells[1], str(unique_doc2[i]) if i < len(unique_doc2 or []) else "")

        doc.add_paragraph()


def add_footer_note(doc, note="Generated by Zilla Parishad AI Document Processor",
                     template_name=None):
    """Add a styled footer note at the end of the document."""
    _add_horizontal_line(doc, COLORS["light_muted"], 0.5)
    _styled_paragraph(
        doc, note,
        font_name="Arial", font_size=8, italic=True,
        color=COLORS["light_muted"],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(4),
    )
